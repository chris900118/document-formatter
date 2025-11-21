#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
公文格式化脚本
使用 python-docx 库处理 Word 文档格式化
"""

import sys
import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def format_document(input_path, output_path, options):
    """
    格式化文档
    
    Args:
        input_path: 输入文件路径
        output_path: 输出文件路径
        options: 格式化选项
            - fontSize: 字体大小（磅）
            - lineSpacing: 行间距（倍数）
            - paragraphSpacing: 段前间距（磅）
            - pageMargins: 页边距（厘米）
    """
    try:
        # 打开文档
        doc = Document(input_path)
        
        # 设置页边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(options['pageMargins']['top'])
            section.bottom_margin = Cm(options['pageMargins']['bottom'])
            section.left_margin = Cm(options['pageMargins']['left'])
            section.right_margin = Cm(options['pageMargins']['right'])
        
        # 设置段落格式
        for paragraph in doc.paragraphs:
            # 设置字体
            for run in paragraph.runs:
                run.font.size = Pt(options['fontSize'])
                run.font.name = '仿宋_GB2312'  # 公文常用字体
                # 设置中文字体
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            
            # 设置行间距
            paragraph_format = paragraph.paragraph_format
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            paragraph_format.line_spacing = options['lineSpacing']
            
            # 设置段前间距
            if options['paragraphSpacing'] > 0:
                paragraph_format.space_before = Pt(options['paragraphSpacing'])
            
            # 设置首行缩进（2个字符）
            if paragraph.text.strip() and not paragraph.style.name.startswith('Heading'):
                paragraph_format.first_line_indent = Cm(0.74)  # 约2个字符
        
        # 设置表格格式（如果有）
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(options['fontSize'])
                            run.font.name = '仿宋_GB2312'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        
        # 保存文档
        doc.save(output_path)
        
        print(json.dumps({
            'success': True,
            'message': '文档格式化成功'
        }))
        
    except Exception as e:
        print(json.dumps({
            'success': False,
            'message': f'格式化失败: {str(e)}'
        }), file=sys.stderr)
        sys.exit(1)


def main():
    """主函数"""
    if len(sys.argv) != 4:
        print(json.dumps({
            'success': False,
            'message': '参数错误：需要 input_path, output_path, options'
        }), file=sys.stderr)
        sys.exit(1)
    
    input_path = sys.argv[1]
    output_path = sys.argv[2]
    options = json.loads(sys.argv[3])
    
    format_document(input_path, output_path, options)


if __name__ == '__main__':
    main()
