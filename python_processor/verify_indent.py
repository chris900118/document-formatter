import os
from docx import Document
from docx.oxml.ns import qn
from formatter import format_document

def verify_indent():
    # 创建一个临时输入文件
    input_path = "temp_input.docx"
    output_path = "temp_output.docx"
    
    doc = Document()
    doc.add_paragraph("这是一个测试段落，用于验证首行缩进。")
    doc.add_paragraph("这是第二个测试段落。")
    doc.save(input_path)
    
    # 定义测试配置
    profile = {
        "styles": {
            "body": {
                "fontFamily": "宋体",
                "fontSize": 12,
                "firstLineIndent": 2  # 设置 2 字符缩进
            }
        }
    }
    
    mappings = {"0": "body", "1": "body"}
    
    # 运行格式化
    result = format_document(input_path, profile, output_path, mappings)
    
    if not result["success"]:
        print(f"格式化失败: {result['error']}")
        return

    # 验证输出文件的 XML
    doc = Document(output_path)
    for i, para in enumerate(doc.paragraphs):
        pPr = para._element.pPr
        if pPr is not None:
            ind = pPr.ind
            if ind is not None:
                firstLineChars = ind.get(qn('w:firstLineChars'))
                firstLine = ind.get(qn('w:firstLine'))
                
                print(f"段落 {i}:")
                print(f"  w:firstLineChars = {firstLineChars}")
                print(f"  w:firstLine = {firstLine}")
                
                if firstLineChars == "200" and firstLine is None:
                    print("  ✓ 验证通过：缩进设置为 2 字符且清除了绝对单位。")
                else:
                    print("  ✗ 验证失败：缩进设置不正确。")
            else:
                print(f"段落 {i}: 找不到 ind 元素")
        else:
            print(f"段落 {i}: 找不到 pPr 元素")

    # 清理
    if os.path.exists(input_path): os.remove(input_path)
    if os.path.exists(output_path): os.remove(output_path)

if __name__ == "__main__":
    verify_indent()
