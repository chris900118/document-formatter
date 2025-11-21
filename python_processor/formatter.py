#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import sys
import os
import json
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from cleaner import ManualNumberingCleaner

# 配置UTF-8输出，避免Windows控制台编码问题（stdout/stderr 同时设置）
if sys.stdout is not None:
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except Exception:
        pass
if sys.stderr is not None:
    try:
        sys.stderr.reconfigure(encoding='utf-8')
    except Exception:
        pass

# --- Numbering Helpers ---

def to_roman(n, lower=False):
    """Convert integer to Roman numeral."""
    if not (0 < n < 4000):
        return str(n)
    val = [
        1000, 900, 500, 400,
        100, 90, 50, 40,
        10, 9, 5, 4,
        1
    ]
    syb = [
        "M", "CM", "D", "CD",
        "C", "XC", "L", "XL",
        "X", "IX", "V", "IV",
        "I"
    ]
    roman_num = ''
    i = 0
    while n > 0:
        for _ in range(n // val[i]):
            roman_num += syb[i]
            n -= val[i]
        i += 1
    return roman_num.lower() if lower else roman_num

def to_chinese_num(n, upper=False):
    """Convert integer to Chinese numeral (simplified)."""
    # Simple implementation for 1-99, can be extended
    chars = ["零", "一", "二", "三", "四", "五", "六", "七", "八", "九"]
    units = ["", "十", "百", "千"]
    if upper:
        chars = ["零", "壹", "贰", "叁", "肆", "伍", "陆", "柒", "捌", "玖"]
        units = ["", "拾", "佰", "仟"]
    
    if n == 0: return chars[0]
    if n == 10: return units[1]
    if 10 < n < 20: return units[1] + chars[n % 10]
    
    s = str(n)
    length = len(s)
    result = ""
    for i, digit in enumerate(s):
        d = int(digit)
        if d != 0:
            result += chars[d] + units[length - i - 1]
        else:
            # Handle zero logic (simplified)
            if i < length - 1 and s[i+1] != '0':
                result += chars[0]
    return result.rstrip(chars[0])

def to_circle_num(n):
    """Convert integer to circled number ①."""
    if 1 <= n <= 20:
        return chr(9311 + n)
    return str(n)

def to_alpha(n, lower=False):
    """Convert integer to A, B, C..."""
    if n < 1: return str(n)
    s = ""
    while n > 0:
        n, remainder = divmod(n - 1, 26)
        s = chr(65 + remainder) + s
    return s.lower() if lower else s

def format_number(n, type_str):
    """Format number based on type string."""
    if type_str == '1': return str(n)
    if type_str == '一': return to_chinese_num(n, False)
    if type_str == '壹': return to_chinese_num(n, True)
    if type_str == '①': return to_circle_num(n)
    if type_str == 'I': return to_roman(n, False)
    if type_str == 'i': return to_roman(n, True)
    if type_str == 'A': return to_alpha(n, False)
    if type_str == 'a': return to_alpha(n, True)
    return str(n)

class NumberingManager:
    def __init__(self, profile):
        self.profile = profile
        self.counters = {
            'heading1': 0,
            'heading2': 0,
            'heading3': 0,
            'heading4': 0
        }
        self.hierarchy = {
            'heading1': None,
            'heading2': 'heading1',
            'heading3': 'heading2',
            'heading4': 'heading3'
        }
        # Cache for the "cascade string" (the parent part of the number)
        # e.g. for H2 "1.1", the cascade string from H1 might be "1"
        self.cascade_cache = {} 

    def reset_children(self, level):
        """Reset counters for children of the given level."""
        levels = ['heading1', 'heading2', 'heading3', 'heading4']
        try:
            idx = levels.index(level)
            for i in range(idx + 1, len(levels)):
                self.counters[levels[i]] = 0
        except ValueError:
            pass

    def get_number_string(self, style_key):
        """Calculate the full number string for the current counter state."""
        styles = self.profile.get('styles', {})
        if style_key not in styles:
            return None
        
        config = styles[style_key].get('numbering')
        if not config or not config.get('enabled'):
            return None

        # 1. Get current value
        val = self.counters.get(style_key, 0)
        type_str = config.get('counterType', '1')
        formatted_val = format_number(val, type_str)
        
        prefix = config.get('prefix', '')
        suffix = config.get('suffix', '')
        
        current_level_str = f"{prefix}{formatted_val}{suffix}"

        # 2. Handle Cascade
        if config.get('cascade'):
            parent_key = self.hierarchy.get(style_key)
            if parent_key:
                parent_cascade_str = self.get_cascade_string(parent_key)
                separator = config.get('separator', '.')
                return f"{parent_cascade_str}{separator}{current_level_str}"
        
        return current_level_str

    def get_cascade_string(self, style_key):
        """
        Get the string that this level contributes to its children.
        This handles the 'Chinese H1 Exception' (第一章 -> 1).
        """
        styles = self.profile.get('styles', {})
        config = styles.get(style_key, {}).get('numbering')
        
        # Default to "1" if something is wrong, but usually we rely on counters
        val = self.counters.get(style_key, 1) # Default to 1 if 0 (shouldn't happen if parent exists)
        if val == 0: val = 1 

        if not config:
            return str(val)

        type_str = config.get('counterType', '1')
        
        # Special Rule: H1 Chinese -> Arabic 1
        if style_key == 'heading1' and type_str in ['一', '壹']:
            return str(val)
        
        # Normal recursion
        formatted_val = format_number(val, type_str)
        prefix = config.get('prefix', '')
        suffix = config.get('suffix', '')
        current_str = f"{prefix}{formatted_val}{suffix}"
        
        if config.get('cascade'):
            parent_key = self.hierarchy.get(style_key)
            if parent_key:
                parent_str = self.get_cascade_string(parent_key)
                separator = config.get('separator', '.')
                return f"{parent_str}{separator}{current_str}"
        
        return current_str

    def process_paragraph(self, para, style_key):
        """Update counters and apply numbering to paragraph."""
        if style_key not in self.counters:
            return

        # Increment counter
        self.counters[style_key] += 1
        # Reset children
        self.reset_children(style_key)
        
        # Generate number string
        num_str = self.get_number_string(style_key)
        
        if num_str:
            # Prepend to paragraph text
            # We need to be careful. If we just prepend text, we might mess up existing runs.
            # Strategy: Insert a new run at the beginning.
            # Also, add a space or tab after the number? Usually a space.
            
            full_text = f"{num_str} "
            
            # Insert at the beginning
            if para.runs:
                # Fix: Directly prepend to the first run's text.
                # Do NOT use add_text() as it appends, and combined with the prepend below, it causes duplication.
                para.runs[0].text = full_text + para.runs[0].text
            else:
                para.add_run(full_text)

def scan_headings(input_path, base_font_size=12):
    """
    扫描Word文档中的标题，智能识别并返回文档结构
    
    Args:
        input_path: 输入Word文档路径
        base_font_size: 基础字号，默认12磅
        
    Returns:
        {
            "success": True/False,
            "structure": [
                {
                    "index": 0,
                    "text": "段落文本",
                    "suggestedStyle": "documentTitle", "heading1", "heading2", "heading3", "heading4", "body",
                    "manual_numbering": {  # 新增字段
                        "type": "arabic" | "chinese" | "parenthesis",
                        "match": "1.1 ",
                        "clean_text": "项目建设目标"
                    }
                },
                ...
            ],
            "error": "错误信息"  # 仅当success=False时存在
        }
    """
    try:
        doc = Document(input_path)
        structure = []
        
        # 初始化编号清洗器
        cleaner = ManualNumberingCleaner()
        
        # 定义样式名称转换函数（移到循环外部）
        def get_display_style_name(style_name_raw):
            """将 Word 内部样式名转换为用户友好的中文名称"""
            if not style_name_raw:
                return '正文'
            style_display_map = {
                'Normal': '正文',
                'List Paragraph': '正文（列表）',
                'Body Text': '正文',
                'Body Text First Indent': '正文（首行缩进）',
                'Body Text First Indent 2': '正文（首行缩进2）',
                'Body Text Indent': '正文（缩进）',
                'Heading 1': '标题 1',
                'Heading 2': '标题 2',
                'Heading 3': '标题 3',
                'Heading 4': '标题 4',
                'Title': '标题',
            }
            # 模糊匹配 Body Text 开头的样式
            if style_name_raw.startswith('Body Text'):
                return style_display_map.get(style_name_raw, '正文（' + style_name_raw.replace('Body Text', '').strip() + '）')
            # 模糊匹配 List Paragraph 开头的样式
            if style_name_raw.startswith('List Paragraph'):
                suffix = style_name_raw.replace('List Paragraph', '').strip()
                return '正文（列表' + (suffix if suffix else '') + '）'
            return style_display_map.get(style_name_raw, style_name_raw)
        
        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            
            # 智能判断样式 - 优先使用 Word 原生样式名称
            suggested_style = "body"  # 默认为正文
            
            # 1. 先检查 Word 样式名称（最可靠）
            style_name = None
            style_id = None
            try:
                if para.style:
                    # 获取样式名称，并清理可能的格式描述
                    raw_name = str(para.style.name) if para.style.name else None
                    if raw_name:
                        # 清洗逻辑：
                        # 1. 如果以"样式 "开头，去掉这个前缀
                        if raw_name.startswith('样式 '):
                            raw_name = raw_name[3:]
                        # 2. 如果包含'+', 取第一部分作为纯样式名
                        if '+' in raw_name:
                            raw_name = raw_name.split('+')[0].strip()
                        # 3. 如果包含':', 取冒号前的部分
                        if ':' in raw_name:
                            raw_name = raw_name.split(':')[0].strip()
                        style_name = raw_name.strip()
                    style_id = para.style.style_id if hasattr(para.style, 'style_id') else None
            except Exception:
                pass
            if style_name:
                # 使用模糊匹配将 Word 样式名映射到我们的样式键
                # 这样可以自动支持所有样式变体，无需硬编码完整列表
                style_lower = style_name.lower()
                
                # 精确匹配优先
                if style_lower == 'title' or style_lower == '标题':
                    suggested_style = 'documentTitle'
                elif style_lower == 'heading 1' or style_lower == '标题 1':
                    suggested_style = 'heading1'
                elif style_lower == 'heading 2' or style_lower == '标题 2':
                    suggested_style = 'heading2'
                elif style_lower == 'heading 3' or style_lower == '标题 3':
                    suggested_style = 'heading3'
                elif style_lower == 'heading 4' or style_lower == '标题 4':
                    suggested_style = 'heading4'
                # 模糊匹配：所有 Body Text、List Paragraph、Normal 及其变体都识别为 body
                elif (style_lower == 'normal' or 
                      style_lower == '正文' or 
                      style_lower.startswith('body text') or 
                      style_lower.startswith('list paragraph') or 
                      style_lower.startswith('列出段落')):
                    suggested_style = 'body'
                
                # 使用转换函数获取显示名称
                display_style_name = get_display_style_name(style_name)
                
                # 如果匹配到已知样式，使用识别的样式键
                if suggested_style != 'body':
                    # 检测手动编号
                    numbering_detection = cleaner.detect(text)
                    
                    item = {
                        "index": idx,
                        "text": text[:100],
                        "suggestedStyle": suggested_style,
                        "suggested_key": suggested_style,
                        "style": display_style_name,  # 使用用户友好的显示名称
                        "styleId": style_id or "",
                        "originalStyleName": style_name  # 保留原始样式名称供调试
                    }
                    
                    # 如果检测到手动编号，添加 manual_numbering 字段
                    if numbering_detection:
                        item["manual_numbering"] = {
                            "type": numbering_detection["type"],
                            "match": numbering_detection["raw_match"],
                            "clean_text": numbering_detection["clean_text"]
                        }
                    
                    structure.append(item)
                    continue
            
            # 2. 如果样式名称无法识别，使用格式推断（兜底逻辑）
            font_size = None
            is_bold = False
            alignment = para.alignment
            
            if para.runs:
                first_run = para.runs[0]
                if first_run.font.size:
                    font_size = first_run.font.size.pt
                if first_run.font.bold:
                    is_bold = True
            
            # 格式推断判断逻辑
            # 1. 文档标题：居中、加粗、字号最大（通常>=16pt）
            if alignment == WD_ALIGN_PARAGRAPH.CENTER and is_bold and font_size and font_size >= 16:
                suggested_style = "documentTitle"
            # 2. 一级标题：加粗、字号较大（14-15pt）
            elif is_bold and font_size and 14 <= font_size < 16:
                suggested_style = "heading1"
            # 3. 二级标题：加粗、字号中等（13pt左右）
            elif is_bold and font_size and font_size == 13:
                suggested_style = "heading2"
            # 4. 三级标题：加粗、字号稍大于正文（12-12.5pt）
            elif is_bold and font_size and 12 <= font_size < 13:
                suggested_style = "heading3"
            # 5. 四级标题：加粗、字号与正文相同
            elif is_bold and font_size and font_size == base_font_size:
                suggested_style = "heading4"
            # 6. 正文：非加粗、普通字号
            else:
                suggested_style = "body"

            # 汇总（使用友好的显示名称）
            if style_name:
                display_name = get_display_style_name(style_name)
            else:
                display_name = style_id or "正文"
            
            # 检测手动编号
            numbering_detection = cleaner.detect(text)
            
            item = {
                "index": idx,
                "text": text[:100],
                "suggestedStyle": suggested_style,
                "suggested_key": suggested_style,
                "style": display_name,
                "styleId": style_id or "",
                "originalStyleName": style_name  # 保留原始样式名称供调试
            }
            
            # 如果检测到手动编号，添加 manual_numbering 字段
            if numbering_detection:
                item["manual_numbering"] = {
                    "type": numbering_detection["type"],
                    "match": numbering_detection["raw_match"],
                    "clean_text": numbering_detection["clean_text"]
                }
            
            structure.append(item)
        
        return {
            "success": True,
            "structure": structure
        }
    
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }

def remove_style_level_numbering(doc):
    """
    移除文档样式定义中的自动编号配置 (w:numPr)
    
    这会清除样式本身携带的编号定义,防止应用样式时自动添加编号
    
    Args:
        doc: python-docx Document 对象
        
    Returns:
        移除的样式编号定义数量
    """
    removed_count = 0
    try:
        # 获取文档样式部件
        styles_part = doc.part.styles
        if not styles_part:
            return 0
        
        # 遍历所有段落样式
        for style in styles_part:
            try:
                # 只处理段落样式
                if not hasattr(style, 'type') or style.type != WD_STYLE_TYPE.PARAGRAPH:
                    continue
                
                # 获取样式的 XML 元素
                style_element = style.element
                
                # 查找 pPr (段落属性)
                pPr = None
                for child in style_element:
                    if child.tag == qn('w:pPr'):
                        pPr = child
                        break
                
                if pPr is None:
                    continue
                
                # 查找并移除所有 numPr 节点
                numPr_nodes = [child for child in pPr if child.tag == qn('w:numPr')]
                if numPr_nodes:
                    for numPr in numPr_nodes:
                        pPr.remove(numPr)
                    removed_count += 1
                    print(f"✓ 样式 [{style.name}] 已移除编号定义", file=sys.stderr)
                    
            except Exception as e:
                print(f"✗ 处理样式失败: {str(e)}", file=sys.stderr)
        
    except Exception as e:
        print(f"✗ 移除样式级编号失败: {str(e)}", file=sys.stderr)
    
    return removed_count

def remove_paragraph_numbering(paragraph):
    """
    移除单个段落的自动编号属性 (w:numPr)
    
    这不会影响用户手打的 "1." 文本,只移除 Word 的自动列表格式
    
    Args:
        paragraph: python-docx Paragraph 对象
        
    Returns:
        dict: {"removed": bool, "level": int or None}
    """
    result = {"removed": False, "level": None}
    
    try:
        # 获取底层 XML 元素
        p_element = paragraph._element
        
        # 查找段落属性 (pPr)
        pPr = None
        for child in p_element:
            if child.tag == qn('w:pPr'):
                pPr = child
                break
        
        if pPr is None:
            return result
        
        # 查找所有 numPr 节点（自动编号属性）
        numPr_nodes = []
        for child in pPr:
            if child.tag == qn('w:numPr'):
                numPr_nodes.append(child)
        
        # 尝试获取编号层级信息（用于调试和潜在的样式推断）
        if numPr_nodes:
            try:
                numPr = numPr_nodes[0]
                for ilvl_node in numPr:
                    if ilvl_node.tag == qn('w:ilvl'):
                        level_val = ilvl_node.get(qn('w:val'))
                        if level_val:
                            result["level"] = int(level_val)
                        break
            except Exception:
                pass
        
        # 移除找到的所有编号节点
        if numPr_nodes:
            for numPr in numPr_nodes:
                pPr.remove(numPr)
            result["removed"] = True
            
    except Exception as e:
        print(f"✗ 移除段落编号失败: {str(e)}", file=sys.stderr)
    
    return result

def format_document(input_path, profile, output_path, mappings=None, text_replacements=None, enable_auto_numbering=True):
    """
    根据配置规范和用户修正后的映射关系格式化Word文档
    
    Args:
        input_path: 输入Word文档路径
        profile: 配置规范字典，包含documentTitle, heading1-4, body的格式定义
        output_path: 输出Word文档路径
        mappings: 用户修正后的映射关系 {段落索引: 样式键}
        text_replacements: 用户修正后的文本内容 {段落索引: 新文本}
        
    Returns:
        {
            "success": True/False,
            "outputPath": "输出路径",
            "error": "错误信息"
        }
    """
    try:
        doc = Document(input_path)

        # 兼容前端 profile 结构：可能为 { styles: {...}, specialRules: {...} }
        styles_dict = profile.get('styles') if isinstance(profile, dict) and 'styles' in profile else profile
        special_rules = profile.get('specialRules', {}) if isinstance(profile, dict) else {}
        
        # 特殊规则：移除自动编号
        # 第一步：移除样式定义中的编号配置（防止应用样式时引入编号）
        if special_rules.get('removeManualNumberPrefixes'):
            print("[DEBUG] 开始移除样式级编号定义...", file=sys.stderr)
            style_removed = remove_style_level_numbering(doc)
            print(f"[DEBUG] 已清理 {style_removed} 个样式的编号定义", file=sys.stderr)
        
        # 初始化编号管理器
        numbering_manager = NumberingManager(profile)

        # 遍历段落应用格式
        for idx, para in enumerate(doc.paragraphs):
            # 0. 优先应用文本替换 (用户纠偏)
            if text_replacements and str(idx) in text_replacements:
                new_text = text_replacements[str(idx)]
                # 如果新文本为空，是否应该删除段落？
                # 目前逻辑：如果为空字符串，则清空段落内容
                # 注意：直接赋值 para.text 会清除所有原有格式(runs)，但对于标题纠偏通常是可以接受的
                para.text = new_text
            
            text = para.text.strip()
            
            # 检测图片段落（特殊规则优先处理）
            has_picture = False
            try:
                for run in para.runs:
                    if run._element.xml and ('drawing' in run._element.xml or 'pict' in run._element.xml):
                        has_picture = True
                        break
            except Exception:
                pass
            
            # 特殊规则：图片单倍行距
            if has_picture and special_rules.get('pictureLineSpacing'):
                try:
                    para.paragraph_format.line_spacing = 1.0
                except Exception:
                    pass
            
            # 特殊规则：图片居中
            if has_picture and special_rules.get('pictureCenterAlign'):
                try:
                    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    pass
            
            # 如果是纯图片段落且无文本，跳过后续文字格式处理
            if has_picture and not text:
                continue
            
            # 获取该段落的样式键
            style_key = mappings.get(str(idx), "body") if mappings else "body"
            # 兼容前端别名：title/normal -> documentTitle/body
            alias_map = { 'title': 'documentTitle', 'normal': 'body' }
            style_key = alias_map.get(style_key, style_key)
            
            # 应用自动编号（受开关控制，且样式中需启用numbering）
            # 注意：这会修改段落文本，必须在后续格式应用之前执行
            if enable_auto_numbering:
                numbering_manager.process_paragraph(para, style_key)
            
            # 特殊规则：第二步 - 移除段落级自动编号属性
            # 这只会移除 Word 的自动列表格式，不会影响用户手打的 "1." 文本
            if special_rules.get('removeManualNumberPrefixes'):
                removal_result = remove_paragraph_numbering(para)
                if removal_result["removed"]:
                    level_info = f" (层级:{removal_result['level']})" if removal_result["level"] is not None else ""
                    print(f"[DEBUG] 段落 {idx} 已移除编号{level_info}: {text[:30]}...", file=sys.stderr)
            
            # 获取对应的样式配置
            if not styles_dict or style_key not in styles_dict:
                continue
            style_config = styles_dict[style_key]
            
            # 应用段落格式
            para_format = para.paragraph_format
            
            # 对齐方式
            alignment_map = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
                "distribute": WD_ALIGN_PARAGRAPH.DISTRIBUTE
            }
            if "alignment" in style_config:
                para_format.alignment = alignment_map.get(style_config["alignment"], WD_ALIGN_PARAGRAPH.LEFT)
            
            # 缩进（单位：字符，需转换为厘米）
            char_width = 0.42  # 每个中文字符约0.42cm（基于四号字）
            if "firstLineIndent" in style_config:
                indent_chars = style_config["firstLineIndent"]
                para_format.first_line_indent = Cm(indent_chars * char_width)
            
            # 行距：数值<=3 视为倍数，否则按磅数
            if "lineSpacing" in style_config and style_config["lineSpacing"] is not None:
                ls = style_config["lineSpacing"]
                try:
                    if isinstance(ls, (int, float)) and ls <= 3:
                        para_format.line_spacing = float(ls)
                    else:
                        para_format.line_spacing = Pt(float(ls))
                except Exception:
                    pass
            
            # 段前段后间距（磅）
            if "spaceBefore" in style_config and style_config["spaceBefore"] is not None:
                try:
                    para_format.space_before = Pt(float(style_config["spaceBefore"]))
                except Exception:
                    pass
            if "spaceAfter" in style_config and style_config["spaceAfter"] is not None:
                try:
                    para_format.space_after = Pt(float(style_config["spaceAfter"]))
                except Exception:
                    pass

            # 特殊规则：重置左右缩进与段前段后间距
            if special_rules.get('resetIndentsAndSpacing'):
                try:
                    para_format.left_indent = Pt(0)
                    para_format.right_indent = Pt(0)
                    para_format.space_before = Pt(0)
                    para_format.space_after = Pt(0)
                except Exception:
                    pass
            
            # 应用字体格式到所有run
            for run in para.runs:
                # 字体
                if "fontFamily" in style_config:
                    font_name = style_config["fontFamily"]
                    run.font.name = font_name
                    # 设置中文/西文字体
                    try:
                        r_pr = run._element.get_or_add_rPr()
                        r_fonts = r_pr.get_or_add_rFonts()
                        r_fonts.set(qn('w:ascii'), font_name)
                        r_fonts.set(qn('w:eastAsia'), font_name)
                        r_fonts.set(qn('w:hAnsi'), font_name)
                        r_fonts.set(qn('w:cs'), font_name)
                    except Exception:
                        pass
                
                # 字号（磅）
                if "fontSize" in style_config:
                    run.font.size = Pt(style_config["fontSize"])
                
                # 加粗
                if "bold" in style_config:
                    run.font.bold = style_config["bold"]
                
                # 颜色（十六进制）
                if "color" in style_config:
                    color_hex = style_config["color"].lstrip('#')
                    r = int(color_hex[0:2], 16)
                    g = int(color_hex[2:4], 16)
                    b = int(color_hex[4:6], 16)
                    run.font.color.rgb = RGBColor(r, g, b)

                # 特殊规则：英文与数字自动 Times New Roman
                if special_rules.get('autoTimesNewRoman'):
                    try:
                        text = run.text or ""
                        has_ascii = any(('A' <= ch <= 'Z') or ('a' <= ch <= 'z') or ('0' <= ch <= '9') for ch in text)
                        if has_ascii:
                            r_pr = run._element.get_or_add_rPr()
                            r_fonts = r_pr.get_or_add_rFonts()
                            r_fonts.set(qn('w:ascii'), 'Times New Roman')
                            r_fonts.set(qn('w:hAnsi'), 'Times New Roman')
                    except Exception:
                        pass
            
            # 应用Word内置样式名（关键：真正改变段落样式属性）
            style_name_map = {
                'documentTitle': 'Title',
                'heading1': 'Heading 1',
                'heading2': 'Heading 2',
                'heading3': 'Heading 3',
                'heading4': 'Heading 4',
                'body': 'Normal'
            }
            if style_key in style_name_map:
                try:
                    para.style = style_name_map[style_key]
                except Exception:
                    pass

            # 再次移除自动编号，避免样式切换引入的编号定义残留
            try:
                p = para._p
                pPr = p.pPr
                if pPr is not None and pPr.numPr is not None:
                    pPr.remove(pPr.numPr)
            except Exception:
                pass
        
        # 特殊规则：第三步 - 保存前最后一次全局清理（确保彻底移除残留编号）
        if special_rules.get('removeManualNumberPrefixes'):
            print("[DEBUG] 执行保存前全局编号清理...", file=sys.stderr)
            removed_count = 0
            checked_count = 0
            detected_count = 0
            
            for para in doc.paragraphs:
                checked_count += 1
                # 检查段落XML结构（调试用）
                p_element = para._element
                has_pPr = any(child.tag == qn('w:pPr') for child in p_element)
                if has_pPr:
                    pPr = next((child for child in p_element if child.tag == qn('w:pPr')), None)
                    if pPr:
                        has_numPr = any(child.tag == qn('w:numPr') for child in pPr)
                        if has_numPr:
                            detected_count += 1
                            print(f"[DEBUG] 发现残留编号: {para.text[:50] if para.text else '(空)'}", file=sys.stderr)
                
                removal_result = remove_paragraph_numbering(para)
                if removal_result["removed"]:
                    removed_count += 1
            
            print(f"[DEBUG] 全局清理完成: 检查 {checked_count} 段落, 检测到 {detected_count} 个编号, 移除 {removed_count} 个", file=sys.stderr)
        
        # 保存文档
        doc.save(output_path)
        
        return {
            "success": True,
            "outputPath": output_path
        }
    
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }

def main():
    """命令行入口"""
    if len(sys.argv) < 2:
        print(json.dumps({"success": False, "error": "缺少命令参数"}, ensure_ascii=False))
        sys.exit(1)
    
    command = sys.argv[1]
    
    if command == "scan_headings":
        if len(sys.argv) < 3:
            print(json.dumps({"success": False, "error": "缺少输入文件路径"}, ensure_ascii=False))
            sys.exit(1)
        
        input_path = sys.argv[2]
        base_font_size = int(sys.argv[3]) if len(sys.argv) > 3 else 12
        
        result = scan_headings(input_path, base_font_size)
        print(json.dumps(result, ensure_ascii=False))
    
    elif command == "format":
        if len(sys.argv) < 5:
            print(json.dumps({"success": False, "error": "缺少必要参数"}, ensure_ascii=False))
            sys.exit(1)
        
        # Electron传递的参数格式: ['format', inputPath, outputPath, JSON.stringify({profile, mappings})]
        input_path = sys.argv[2]
        output_path = sys.argv[3]
        payload_str = sys.argv[4]
        payload = json.loads(payload_str)
        
        profile = payload.get("profile")
        mappings = payload.get("mappings", {})
        text_replacements = payload.get("text_replacements", {})
        enable_auto_numbering = payload.get("enable_auto_numbering", True)
        
        if not all([input_path, profile, output_path]):
            print(json.dumps({"success": False, "error": "缺少必要参数"}, ensure_ascii=False))
            sys.exit(1)
        
        result = format_document(input_path, profile, output_path, mappings, text_replacements, enable_auto_numbering)
        print(json.dumps(result, ensure_ascii=False))
    
    else:
        print(json.dumps({"success": False, "error": f"未知命令: {command}"}, ensure_ascii=False))
        sys.exit(1)

if __name__ == "__main__":
    main()
