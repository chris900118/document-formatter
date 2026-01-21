import os
from docx import Document
from formatter import format_document

def verify_spacing():
    input_path = "temp_input_spacing.docx"
    output_path = "temp_output_spacing.docx"
    
    doc = Document()
    doc.add_paragraph("测试内容")
    doc.save(input_path)
    
    profile = {
        "styles": {
            "body": {
                "fontFamily": "仿宋",
                "fontSize": 16,
                "lineSpacing": 28,
            }
        }
    }
    mappings = {"0": "body"}
    
    format_document(input_path, profile, output_path, mappings)
    
    doc = Document(output_path)
    for i, para in enumerate(doc.paragraphs):
        print(f"--- Paragraph {i} XML ---")
        # 直接使用 .xml 属性获取 XML 文本
        print(para._element.xml)

    if os.path.exists(input_path): os.remove(input_path)
    if os.path.exists(output_path): os.remove(output_path)

if __name__ == "__main__":
    verify_spacing()
